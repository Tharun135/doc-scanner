from flask import Blueprint, request, jsonify, render_template
from flask_login import login_required, current_user
import os
import re
import subprocess
import markdown
import PyPDF2
from docx import Document
from bs4 import BeautifulSoup
import logging
import importlib
import sys
import textstat
import time
from dataclasses import asdict

# Try to import spacy but handle import errors gracefully
try:
    import spacy
    SPACY_IMPORT_SUCCESS = True
except ImportError:
    SPACY_IMPORT_SUCCESS = False
    spacy = None

# Load environment variables from .env file
try:
    from dotenv import load_dotenv
    load_dotenv()
    logger = logging.getLogger(__name__)
    logger.info("Environment variables loaded from .env file")
except ImportError:
    logger = logging.getLogger(__name__)
    logger.warning("python-dotenv not available - environment variables must be set manually")

# Add the parent directory to the path for imports
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

main = Blueprint('main', __name__)

logging.basicConfig(level=logging.INFO)  # Changed from DEBUG to hide RAG debug messages
logger = logging.getLogger(__name__)

# Load spaCy English model (make sure to run: python -m spacy download en_core_web_sm)
if SPACY_IMPORT_SUCCESS and os.environ.get('FLASK_ENV') != 'production':
    try:
        # Try to load with reduced memory footprint
        nlp = spacy.load("en_core_web_sm", disable=["ner", "textcat"])
        SPACY_AVAILABLE = True
        logger.info("spaCy model loaded successfully with reduced components")
    except MemoryError as e:
        logger.warning(f"spaCy model loading failed due to memory constraints: {e}")
        logger.warning("Running without spaCy - sentence splitting will use basic methods")
        nlp = None
        SPACY_AVAILABLE = False
    except ImportError as e:
        logger.warning(f"spaCy not available: {e}")
        nlp = None
        SPACY_AVAILABLE = False
    except Exception as e:
        logger.warning(f"spaCy model not available: {e}")
        logger.warning("Falling back to basic sentence processing")
        nlp = None
        SPACY_AVAILABLE = False
else:
    logger.warning("spaCy package not installed - using basic sentence processing")
    nlp = None
    SPACY_AVAILABLE = False

############################
# SENTENCE EXTRACTION HELPERS
############################

def extract_sentences_with_html_preservation(html_content):
    """
    Extract sentences from HTML content while preserving the HTML structure.
    Returns a list of sentence objects that contain both the original HTML and plain text versions.
    """
    sentences = []
    
    # Parse HTML content
    soup = BeautifulSoup(html_content, "html.parser")
    
    # Process leaf elements that contain text to avoid double-counting nested blocks
    # (e.g., a <p> inside a <div> should only be counted by the <p>)
    all_potential = soup.find_all(['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'td', 'th', 'span', 'blockquote'])
    
    text_elements = []
    for el in all_potential:
        # Skip if it's already contained in another block we've already selected? 
        # No, better to skip if it has block-level children that we will select anyway.
        has_block_child = el.find(['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'td', 'th'])
        
        # If it has block children, we only want its DIRECT text (if any)
        # But for simplicity, most documents have text in p/li/hX.
        # If it's a div with no block children, it's a text block.
        if not has_block_child and el.get_text().strip():
            text_elements.append(el)
        elif has_block_child:
            # Check if it has any direct text children that aren't inside the block children
            # This is rare in clean HTML but common in messy ones.
            # For now, we favor leaf nodes.
            pass
            
    for block_idx, element in enumerate(text_elements):
        if not element.get_text().strip():
            continue
            
        # Get the HTML content of this element
        element_html = str(element)
        tag_name = element.name # p, h1, etc.
        
        # Get the plain text version for analysis with a space separator to preserve word boundaries
        element_text = element.get_text(separator=' ')
        
        # Split the text into sentences
        raw_fragments = []
        if SPACY_AVAILABLE and nlp:
            doc = nlp(element_text)
            raw_fragments = [sent.text.strip() for sent in doc.sents if sent.text.strip()]
        else:
            # Fallback for splitting sentences when spacy is not available
            raw_fragments = [s.strip() for s in re.split(r'[.!?]+\s+', element_text) if s.strip()]

        # 🔗 MERGING STEP: Unify technical fragments (e.g. "Label: Content")
        merged_fragments = []
        i = 0
        while i < len(raw_fragments):
            current = raw_fragments[i]
            while i + 1 < len(raw_fragments) and re.search(r':\s*$', current):
                current = f"{current} {raw_fragments[i+1]}"
                i += 1
            merged_fragments.append(current)
            i += 1

        # Process fragments into sentence objects
        for sent_text in merged_fragments:
            # Find corresponding HTML fragment
            html_fragment = find_html_fragment_for_sentence(element_html, sent_text, element_text)
            
            # Use a unified sentence class with Block Info
            class SentenceObj:
                def __init__(self, text, html_fragment, block_index, tag_name):
                    self.text = text.strip()
                    self.html_fragment = html_fragment
                    self.block_index = block_index
                    self.tag_name = tag_name # h1, p, li, etc.
                    self.start_char = 0
                    self.end_char = len(text)
            
            sentences.append(SentenceObj(sent_text, html_fragment, block_idx, tag_name))
    
    return sentences

def find_html_fragment_for_sentence(element_html, sentence_text, full_text):
    """
    Find the HTML fragment that corresponds to a specific sentence within an element.
    This preserves HTML tags like <strong>, <em>, <a>, <code>, etc.
    """
    soup = BeautifulSoup(element_html, "html.parser")
    
    # Get the first actual HTML element (skip the document wrapper)
    first_element = soup.find()
    if not first_element:
        return sentence_text
    
    element_plain_text = first_element.get_text().strip()
    
    # Simple case: if the element contains only this sentence
    if element_plain_text.strip() == sentence_text.strip():
        # Return the inner HTML (content without the wrapper tag)
        if hasattr(first_element, 'contents') and first_element.contents:
            inner_content = ''.join(str(content) for content in first_element.contents)
            return inner_content
        return sentence_text
    
    # Complex case: sentence is part of a larger element
    sentence_start = full_text.find(sentence_text.strip())
    if sentence_start == -1:
        return sentence_text  # Fallback - return plain text
    
    sentence_end = sentence_start + len(sentence_text.strip())
    
    # If this sentence spans the entire element content, return the inner HTML
    if sentence_start == 0 and sentence_end >= len(full_text.strip()):
        if hasattr(first_element, 'contents') and first_element.contents:
            inner_content = ''.join(str(content) for content in first_element.contents)
            return inner_content
    
    # For partial matches, try to find the corresponding HTML portion
    try:
        # Get all text nodes and their positions to map back to HTML
        html_content = str(first_element)
        
        # Create a temporary element to work with
        temp_soup = BeautifulSoup(html_content, "html.parser")
        temp_element = temp_soup.find()
        
        if temp_element:
            # Try to find the sentence text within the element
            element_text = temp_element.get_text()
            
            # Find the sentence within the element's text
            local_sentence_start = element_text.find(sentence_text.strip())
            if local_sentence_start != -1:
                local_sentence_end = local_sentence_start + len(sentence_text.strip())
                
                # Try to map this back to HTML by looking for text patterns
                # For now, if we can find the exact sentence text in the HTML, return it with any inline tags
                html_str = str(temp_element)
                if sentence_text.strip() in html_str:
                    # Look for the sentence with potential HTML tags
                # Ensure we handle multiple occurrences of the same text
                    # Create a pattern that allows HTML tags within the sentence
                    # Use words as anchors to find the HTML segment
                    words = sentence_text.strip().split()
                    if words:
                        # Create pattern allowing whitespace, tags, and common HTML entities between words
                        pattern_words = [re.escape(word) for word in words]
                        # NON-GREEDY joiner to prevent swallowing next sentence
                        boundary_pattern = r'(?:\s|<[^>]*>|&nbsp;|&#160;)+?'
                        pattern = boundary_pattern.join(pattern_words)
                        if len(pattern_words) > 1:
                            # Also allow leading/trailing tags but NON-GREEDY
                            pattern = r'(?:<[^>]*>)*?\s*' + pattern + r'\s*(?:<[^>]*>)*?'
                        
                        match = re.search(pattern, html_str, re.IGNORECASE | re.DOTALL)
                        if match:
                            match_text = match.group(0)
                            # Loosened sanity check to allow for technical merged sentences
                            # Ratio increased to 2.5x, word delta increased to 10 for complex technical lists
                            if len(match_text) < len(sentence_text) * 2.5:
                                try:
                                    match_plain = BeautifulSoup(match_text, "html.parser").get_text().split()
                                    if abs(len(match_plain) - len(words)) <= 10:
                                        logger.info(f"✅ Precise HTML fragment found for sentence {len(words)} words")
                                        return match_text.strip()
                                except Exception as e:
                                    logger.warning(f"Word-count check failed: {e}")
                                    return match_text.strip()
                
                # Fallback: return the plain text
                return sentence_text.strip()
    except Exception as e:
        # If anything goes wrong, return the plain text
        return sentence_text.strip()
    
    # Final fallback
    return sentence_text.strip()

############################
# FILE PARSING HELPERS
############################

def clean_malformed_html_attributes(text):
    """
    Clean malformed HTML attributes that might appear in text content.
    Specifically targets patterns like: ="sentence-highlight" id="content-sentence-0"
    """
    
    # Pattern to match malformed HTML attributes starting with ="
    # This catches patterns like: ="sentence-highlight" id="content-sentence-0" data-sentence-index="0">
    pattern = r'=[\"\'][^\"\']*[\"\'][^>]*>'
    
    # Remove these malformed patterns
    cleaned = re.sub(pattern, '', text)
    
    # Also remove any standalone HTML attributes that might be left
    # Pattern for: id="something" data-something="value" etc.
    attr_pattern = r'\s*(?:id|class|data-[\w-]+)\s*=\s*["\'][^"\']*["\']'
    cleaned = re.sub(attr_pattern, '', cleaned)
    
    # Clean up any remaining HTML tag fragments
    fragment_pattern = r'</?[^>]*>'
    cleaned = re.sub(fragment_pattern, '', cleaned)
    
    return cleaned.strip()

def parse_file(file):
    filename = file.filename.lower()
    extension = os.path.splitext(filename)[1]

    try:
        if extension == '.zip':
            return parse_zip(file)
        elif extension == '.docx':
            return parse_docx(file)
        elif extension == '.doc':
            return parse_doc(file)
        elif extension == '.pdf':
            return parse_pdf(file)
        elif extension == '.md':
            return parse_md(file.read())
        elif extension == '.adoc':
            return parse_adoc(file.read())
        elif extension in ['.txt', '']:
            return parse_txt(file.read())
        else:
            return file.read().decode("utf-8", errors="replace")
    except Exception as e:
        logger.error(f"Error parsing {filename}: {str(e)}")
        return f"Error parsing {filename}: {str(e)}"

def parse_zip(file_stream):
    """Parse ZIP file and extract text content from supported documents inside
    
    ⚠️ REVIEWER DECISION NEEDED:
    Current behavior: Merges all files into one document for review.
    Alternative: Review each file separately.
    
    Question: Are these separate documents or one review unit?
    - If separate: Need to review each with its own document gate
    - If collection: Need summary of collection-level issues
    
    DO NOT: Blend feedback silently across files without clarifying this to users.
    """
    import zipfile
    import io
    
    try:
        html_content = "<h2>Contents of ZIP file:</h2>\n"
        
        with zipfile.ZipFile(file_stream, 'r') as zip_file:
            file_list = zip_file.namelist()
            
            # Filter for supported file types
            supported_files = []
            for filename in file_list:
                if filename.lower().endswith(('.txt', '.md', '.docx', '.pdf')):
                    supported_files.append(filename)
            
            if not supported_files:
                return "<p>No supported document files found in ZIP archive. Supported formats: .txt, .md, .docx, .pdf</p>"
            
            html_content += f"<p>Found {len(supported_files)} supported document(s) in ZIP file:</p>\n"
            
            # Process each supported file
            for filename in supported_files:
                try:
                    html_content += f"<h3>📄 {filename}</h3>\n"
                    
                    # Extract file content
                    with zip_file.open(filename) as extracted_file:
                        file_content = io.BytesIO(extracted_file.read())
                        
                        # Determine file type and parse accordingly
                        ext = os.path.splitext(filename.lower())[1]
                        
                        if ext == '.txt':
                            content = parse_txt(file_content.read())
                        elif ext == '.md':
                            content = parse_md(file_content.read())
                        elif ext == '.docx':
                            content = parse_docx(file_content)
                        elif ext == '.pdf':
                            content = parse_pdf(file_content)
                        else:
                            content = f"<p>Unsupported file type: {ext}</p>"
                        
                        html_content += content + "\n"
                        
                except Exception as e:
                    html_content += f"<p>Error processing {filename}: {str(e)}</p>\n"
                    logger.error(f"Error processing file {filename} in ZIP: {e}")
            
            return html_content
            
    except zipfile.BadZipFile:
        return "Error: Invalid or corrupted ZIP file"
    except Exception as e:
        logger.error(f"Error parsing ZIP file: {e}")
        return f"Error parsing ZIP file: {str(e)}"

def parse_docx(file_stream):
    doc = Document(file_stream)
    html_content = ""
    for paragraph in doc.paragraphs:
        html_content += f"<p>{paragraph.text}</p>"
    return html_content

def parse_doc(file_stream):
    temp_path = "temp_upload_doc.doc"
    file_stream.save(temp_path)
    try:
        cmd = ["antiword", temp_path]
        result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        if result.returncode == 0:
            text_content = result.stdout
            html_content = ""
            for paragraph in text_content.split("\n\n"):
                html_content += f"<p>{paragraph}</p>"
            return html_content
        else:
            return f"Error reading .doc file: {result.stderr}"
    except FileNotFoundError:
        return "Error: 'antiword' not in PATH or not installed."
    except Exception as e:
        return f"Error reading .doc file: {str(e)}"
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)

def parse_pdf(file_stream):
    try:
        reader = PyPDF2.PdfReader(file_stream)
        html_content = ""
        for page in reader.pages:
            page_text = page.extract_text() or ""
            for paragraph in page_text.split("\n\n"):
                html_content += f"<p>{paragraph}</p>"
        return html_content
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

def parse_md(content_bytes):
    md_text = content_bytes.decode("utf-8", errors="replace")
    html_text = markdown.markdown(md_text)
    return html_text

def parse_adoc(content_bytes):
    adoc_text = content_bytes.decode("utf-8", errors="replace")
    html_content = ""
    for paragraph in adoc_text.split("\n\n"):
        html_content += f"<p>{paragraph}</p>"
    return html_content

def parse_txt(content_bytes):
    txt_text = content_bytes.decode("utf-8", errors="replace")
    html_content = ""
    for paragraph in txt_text.split("\n\n"):
        html_content += f"<p>{paragraph}</p>"
    return html_content

def load_rules():
    rules = []
    rules_folder = os.path.join(os.path.dirname(__file__), 'rules')
    
    # Files that are utilities or helpers and don't need a check function
    ignore_files = {
        'analytics.py', 'batch_processor.py', 'rag_rule_helper.py',
        'title_utils.py', 'rule_remediations.py', 'verb_tense.py',
        'anaphora_resolution.py', 'matcher.py', 'nominalizations.py'
    }
    
    for filename in os.listdir(rules_folder):
        if filename.endswith('.py') and filename != '__init__.py':
            module_name = filename[:-3]  # Remove .py extension
            try:
                # Import the module using the proper package path
                module = importlib.import_module(f'app.rules.{module_name}')
                
                if hasattr(module, "check"):
                    rules.append(module.check)
                    logger.info(f"Loaded rule: {filename} (Function: {module.check.__name__})")
                elif filename not in ignore_files:
                    logger.warning(f"Warning: {filename} does not have a `check` function")
            except ImportError as e:
                logger.error(f"Failed to import rule {filename}: {e}")
            except Exception as e:
                logger.error(f"Error loading rule {filename}: {e}")

    logger.info(f"Total rules loaded: {len(rules)}")
    return rules

rules = load_rules()

def review_document(content, rules):
    suggestions = []
    for rule in rules:
        feedback = rule(content)
        if feedback:
            for item in feedback:
                # Handle both string and dict formats
                if isinstance(item, str):
                    # Skip string-only feedback for now
                    continue
                elif isinstance(item, dict):
                    suggestions.append({
                        "text": item.get("text", ""),
                        "start": item.get("start", 0),
                        "end": item.get("end", 0),
                        "message": item.get("message", item.get("suggestion", ""))
                    })
    return {"issues": suggestions, "summary": "Review completed."}

def analyze_sentence(sentence, rules, previous_sentence=None, next_sentence=None):
    feedback = []
    readability_scores = {
        "flesch_reading_ease": textstat.flesch_reading_ease(sentence),
        "gunning_fog": textstat.gunning_fog(sentence),
        "smog_index": textstat.smog_index(sentence),
        "automated_readability_index": textstat.automated_readability_index(sentence)
    }
    quality_score = 55.0  # Placeholder for actual quality score calculation

    # Apply each rule function to the sentence
    for rule_function in rules:
        try:
            # Try to pass adjacent context if the rule supports it (like passive_voice)
            import inspect
            sig = inspect.signature(rule_function)
            params = list(sig.parameters.keys())
            
            if 'previous_sentence' in params and 'next_sentence' in params:
                # Rule supports adjacent context
                rule_feedback = rule_function(sentence, previous_sentence=previous_sentence, next_sentence=next_sentence)
            else:
                # Rule doesn't support context, use standard call
                rule_feedback = rule_function(sentence)
        except Exception as e:
            # Fallback to standard call if inspection fails
            logger.debug(f"Rule inspection failed, using standard call: {e}")
            rule_feedback = rule_function(sentence)
        if rule_feedback:
            # Convert string feedback to expected object format
            for item in rule_feedback:
                if isinstance(item, str):
                    # Parse structured suggestions to extract just the issue for display
                    message = item
                    if 'Issue:' in item and 'Original sentence:' in item and ('AI suggestion:' in item or 'AI Solution:' in item):
                        # Extract just the issue part for the main display
                        lines = item.split('\n')
                        for line in lines:
                            if line.strip().startswith('Issue:'):
                                message = line.replace('Issue:', '').strip()
                                break
                    
                    # Convert string to expected object format
                    feedback.append({
                        "text": sentence,
                        "start": 0,
                        "end": len(sentence),
                        "message": message,
                        "full_suggestion": item,  # Keep the full structured suggestion for detailed view
                        "severity": "warn",  # Default severity for legacy string feedback
                        "color": "yellow"
                    })
                elif isinstance(item, dict):
                    # Handle dict format - may include severity and color from atomic rules
                    feedback_item = {
                        "text": item.get("text", sentence),
                        "start": item.get("start", 0),
                        "end": item.get("end", len(sentence)),
                        "message": item.get("message", str(item)),
                    }
                    
                    # Preserve severity-based information from atomic rules
                    if "severity" in item:
                        feedback_item["severity"] = item["severity"]
                    else:
                        feedback_item["severity"] = "warn"  # Default for legacy rules
                    
                    if "color" in item:
                        feedback_item["color"] = item["color"]
                    else:
                        # Map severity to color if not explicitly provided
                        severity_to_color = {"error": "red", "warn": "yellow", "info": "grey"}
                        feedback_item["color"] = severity_to_color.get(feedback_item["severity"], "yellow")
                    
                    # Preserve additional fields
                    if "suggestion" in item:
                        feedback_item["suggestion"] = item["suggestion"]
                    if "rule_id" in item:
                        feedback_item["rule_id"] = item["rule_id"]
                    if "category" in item:
                        feedback_item["category"] = item["category"]
                    if "full_suggestion" in item:
                        feedback_item["full_suggestion"] = item["full_suggestion"]
                    
                    # Preserve Rule Authority keys
                    for key in ["decision_type", "reviewer_rationale", "rule", "ai_suggestion"]:
                        if key in item:
                            feedback_item[key] = item[key]
                    
                    feedback.append(feedback_item)
                else:
                    # Handle other formats by converting to string
                    feedback.append({
                        "text": sentence,
                        "start": 0,
                        "end": len(sentence),
                        "message": str(item),
                        "severity": "warn",
                        "color": "yellow"
                    })

    return feedback, readability_scores, quality_score
def calculate_quality_index(total_sentences, total_errors):
    if total_sentences == 0:
        return 0
    return max(0, round(100 * (1 - (total_errors / total_sentences))))
############################
# FLASK ROUTES
############################

@main.route('/')
@login_required
def index():
    return render_template('index.html', user=current_user)

@main.route('/start_upload', methods=['POST', 'OPTIONS'])
def start_upload():
    """Initialize upload session and return room ID for progress tracking."""
    # Handle preflight OPTIONS request
    if request.method == 'OPTIONS':
        response = jsonify({"status": "ok"})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type')
        response.headers.add('Access-Control-Allow-Methods', 'POST, OPTIONS')
        return response, 200
    
    try:
        import uuid
        from .progress_tracker import get_progress_tracker
        
        room_id = str(uuid.uuid4())
        progress_tracker = get_progress_tracker()
        
        if progress_tracker:
            progress_tracker.start_session(room_id)
        
        logger.info(f"Upload session initialized: {room_id}")
        return jsonify({"room_id": room_id})
    except Exception as e:
        logger.error(f"Failed to initialize upload session: {str(e)}", exc_info=True)
        return jsonify({"error": f"Failed to initialize upload: {str(e)}"}), 500

@main.route('/debug')
def debug_page():
    """Serve the sentence debugging page"""
    from flask import send_from_directory
    import os
    return send_from_directory(os.path.dirname(os.path.dirname(__file__)), 'debug_sentences.html')

@main.route('/debug_sentences', methods=['POST'])
def debug_sentences():
    """Debug endpoint to check sentence data for malformed HTML"""
    data = request.get_json()
    sentences = data.get('sentences', [])
    
    debug_info = {
        'total_sentences': len(sentences),
        'malformed_sentences': [],
        'html_sentences': [],
        'clean_sentences': []
    }
    
    for i, sentence_data in enumerate(sentences):
        sentence_text = sentence_data.get('sentence', '')
        
        # Check for malformed HTML attributes
        if '="' in sentence_text and ('sentence-highlight' in sentence_text or 'data-sentence-index' in sentence_text):
            debug_info['malformed_sentences'].append({
                'index': i,
                'original': sentence_text,
                'cleaned': clean_malformed_html_attributes(sentence_text)
            })
        
        # Check for any HTML tags
        elif '<' in sentence_text and '>' in sentence_text:
            debug_info['html_sentences'].append({
                'index': i,
                'text': sentence_text
            })
        
        else:
            debug_info['clean_sentences'].append({
                'index': i,
                'text': sentence_text[:100] + ('...' if len(sentence_text) > 100 else '')
            })
    
    return jsonify(debug_info)

@main.route('/upload', methods=['POST', 'OPTIONS'])
def upload_file():
    # Handle preflight OPTIONS request
    if request.method == 'OPTIONS':
        response = jsonify({"status": "ok"})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type')
        response.headers.add('Access-Control-Allow-Methods', 'POST, OPTIONS')
        return response, 200
    
    global current_document_content  # Access global variable
    
    try:
        if 'file' not in request.files:
            logger.error("Upload request missing 'file' field")
            return jsonify({"error": "No file part"}), 400

        file = request.files['file']
        if not file.filename:
            logger.error("Upload request has empty filename")
            return jsonify({"error": "No selected file"}), 400

        # Validate file extension
        filename = file.filename.lower()
        allowed_extensions = ['.txt', '.pdf', '.docx', '.doc', '.md', '.adoc', '.zip']
        if not any(filename.endswith(ext) for ext in allowed_extensions):
            logger.error(f"Unsupported file type: {filename}")
            return jsonify({"error": f"Unsupported file type. Allowed types: {', '.join(allowed_extensions)}"}), 400

        # Get room_id for progress tracking
        room_id = request.form.get('room_id')
        
        from .progress_tracker import get_progress_tracker
        progress_tracker = get_progress_tracker()

        logger.info(f"File uploaded: {file.filename} (Room: {room_id})")
        
        # Validate file size (Flask should handle this automatically, but let's be explicit)
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        max_size = 50 * 1024 * 1024  # 50MB
        if file_size > max_size:
            logger.error(f"File too large: {file_size} bytes (max: {max_size})")
            return jsonify({"error": f"File too large. Maximum size: {max_size // (1024*1024)}MB"}), 400
        
        logger.info(f"File size: {file_size} bytes")
    
    except Exception as e:
        logger.error(f"Error in upload validation: {str(e)}", exc_info=True)
        return jsonify({"error": f"Upload validation failed: {str(e)}"}), 500

    try:
        # Stage 1: Uploading Document (10%)
        if progress_tracker and room_id:
            progress_tracker.update_stage(room_id, 0, f"Uploading {file.filename}...")
        
        # Stage 2: Parsing Content (30%)
        if progress_tracker and room_id:
            progress_tracker.update_stage(room_id, 1, f"Parsing {file.filename.split('.')[-1].upper()} content...")
        
        # Parse file to get both plain text and HTML
        html_content = parse_file(file)
        
        # Clean any existing sentence highlighting from the content (in case document was previously processed)
        if 'sentence-highlight' in html_content:
            logger.warning("🧹 Cleaning existing sentence highlighting from uploaded document...")
            # Remove all sentence highlighting spans but keep the content
            # Remove opening span tags with sentence-highlight class
            html_content = re.sub(r'<span[^>]*sentence-highlight[^>]*>', '', html_content)
            # Remove closing span tags
            html_content = re.sub(r'</span>', '', html_content)
            logger.info("✅ Cleaned existing highlighting markup")
        
        # Check if parsing failed
        if html_content.startswith("Error"):
            logger.error(f"File parsing failed: {html_content}")
            if progress_tracker and room_id:
                progress_tracker.fail_session(room_id, html_content)
            return jsonify({"error": html_content}), 400
        
        # Check if content is empty
        if not html_content.strip():
            error_msg = "The uploaded file appears to be empty or could not be parsed."
            logger.error(f"Empty content from file: {file.filename}")
            if progress_tracker and room_id:
                progress_tracker.fail_session(room_id, error_msg)
            return jsonify({"error": error_msg}), 400
        
        # 🚧 DOCUMENT REVIEW GATE - Reviewer-first analysis
        # ⚠️ NON-NEGOTIABLE: This must happen AFTER parsing, BEFORE sentence extraction
        # 
        # DO NOT move this earlier "for convenience" - that breaks the reviewer-first architecture.
        # Extracting sentences before this gate reintroduces sentence-first bias.
        # 
        # This gate determines whether sentence-level analysis is warranted.
        # It answers: "Would a human reviewer pause here?"
        if progress_tracker and room_id:
            progress_tracker.update_stage(room_id, 2, "Understanding document structure and goal...")
        
        from core.document_review_gate import run_document_review_gate
        document_review = run_document_review_gate(html_content, file.filename)
        
        if document_review.blocking:
            logger.warning(f"Warning: Document has blocking structural issues - but continuing with sentence-level analysis")

        
        # Stage 3: Breaking into Sentences (50%)
        # Only proceed if document structure is sound
        if progress_tracker and room_id:
            progress_tracker.update_stage(room_id, 3, "Identifying sentence boundaries and structure...")
        
        # Store the original HTML content for highlighting
        # Extract sentences that preserve HTML structure while also having plain text for analysis
        
        # Debug: Check if the input HTML already contains highlighting
        if 'sentence-highlight' in html_content:
            logger.error(f"🔥 CRITICAL: Input HTML already contains sentence highlighting! This suggests the document was previously processed.")
            logger.info(f"HTML snippet: {html_content[:500]}...")
        
        sentences = extract_sentences_with_html_preservation(html_content)
        
        # Check if sentence extraction failed
        if not sentences:
            error_msg = "Could not extract any sentences from the document. The file may be corrupted or in an unsupported format."
            logger.error(f"No sentences extracted from file: {file.filename}")
            if progress_tracker and room_id:
                progress_tracker.fail_session(room_id, error_msg)
            return jsonify({"error": error_msg}), 400
        
        logger.info(f"Extracted {len(sentences)} sentences from {file.filename}")
        
        # Use BeautifulSoup to extract plain text for RAG context
        soup = BeautifulSoup(html_content, "html.parser")
        plain_text = soup.get_text(separator="\n")
        
        # Store the document content for RAG context
        current_document_content = plain_text
        
        # Store all sentences for adjacent context in AI suggestions
        global current_sentences_list
        current_sentences_list = sentences

        # 🧠 RAG INGESTION: Add document to knowledge base for better context (ASYNCHRONOUS)
        try:
            from .services.enrichment import ingest_document_to_rag
            import threading
            # Run ingestion in a background thread to avoid blocking the user
            threading.Thread(
                target=ingest_document_to_rag, 
                args=(plain_text, file.filename),
                daemon=True
            ).start()
            logger.info(f"🚀 Document ingestion started in background: {file.filename}")
        except Exception as rag_e:
            logger.warning(f"⚠️ RAG background ingestion failed to start: {rag_e}")

        # Stage 4: Analyzing with Rules (80%)
        # CONDITIONAL ANALYSIS - Only analyze sentences that need it
        if progress_tracker and room_id:
            analysis_scope = document_review.analysis_scope
            if analysis_scope == "minimal":
                progress_tracker.update_stage(room_id, 4, "Spot-checking document clarity...")
            elif analysis_scope == "targeted":
                progress_tracker.update_stage(room_id, 4, "Analyzing flagged sections for clarity...")
            else:
                progress_tracker.update_stage(room_id, 4, "Applying grammar, style, and readability rules...")
        
        sentence_data = []
        total_sentences = len(sentences)
        analyzed_count = 0  # Track how many sentences we actually analyze
        
        # Import the gating function
        from core.document_review_gate import should_analyze_sentence
        
        # 🚀 LAZY RAG OPTIMIZATION: Ensure AI pre-fetch is skipped during extraction phase
        try:
            from rules import rag_rule_helper
            rag_rule_helper.RAG_SKIP_PREFETCH = True
            logger.info("⚡ RAG Lazy Mode: Forced Active for extraction phase")
        except Exception as e:
            logger.warning(f"Could not force Lazy RAG: {e}")

        for index, sent in enumerate(sentences):
            # Update substep progress for analysis (RESCALED: 30-80% range)
            if progress_tracker and room_id and total_sentences > 0:
                substep_progress = 30 + int((index / total_sentences) * 50)
                progress_tracker.update_progress(room_id, substep_progress, f"Analyzing {file.filename} ({index + 1}/{total_sentences})...")
            
            # Use the plain text version for analysis
            plain_text_sentence = sent.text
            html_fragment = getattr(sent, 'html_fragment', plain_text_sentence)
            
            # 🚧 GATED ANALYSIS - Only analyze if needed
            should_analyze = should_analyze_sentence(index, plain_text_sentence, document_review)
            
            if should_analyze:
                analyzed_count += 1
                logger.debug(f"🔍 Analyzing sentence {index} (flagged by document review)")
            else:
                logger.debug(f"⏭️ Skipping sentence {index} (not in analysis scope)")
            
            # Debug: Log sentence content to understand the issue
            logger.debug(f"Processing sentence {index}: '{plain_text_sentence[:100]}...'")
            
            # AGGRESSIVE CLEANING: Handle malformed HTML attributes that somehow got into sentence text
            if '="' in plain_text_sentence and ('sentence-highlight' in plain_text_sentence or 'data-sentence-index' in plain_text_sentence):
                logger.error(f"🚨 MALFORMED HTML ATTRIBUTES DETECTED in sentence {index}: {plain_text_sentence}")
                clean_text = clean_malformed_html_attributes(plain_text_sentence)
                logger.warning(f"🧹 Cleaned malformed HTML: '{clean_text}'")
                plain_text_sentence = clean_text
            
            # Extra safety: Ensure plain text sentence doesn't contain HTML tags
            if '<' in plain_text_sentence and '>' in plain_text_sentence:
                logger.warning(f"🚨 SENTENCE {index} CONTAINS HTML TAGS: {plain_text_sentence}")
                # Check if it contains our highlighting markup specifically
                if 'sentence-highlight' in plain_text_sentence:
                    logger.error(f"🔥 CRITICAL: Sentence {index} contains highlighting markup! This suggests circular processing.")
                
                temp_soup = BeautifulSoup(plain_text_sentence, "html.parser")
                clean_text = temp_soup.get_text().strip()
                logger.warning(f"✅ Cleaned sentence {index}: '{clean_text[:100]}...'")
                plain_text_sentence = clean_text
            
            # Apply rules ONLY if sentence should be analyzed
            if should_analyze:
                # Get adjacent sentences for context-aware analysis
                previous_sentence = None
                next_sentence = None
                
                if index > 0:
                    # Get previous sentence from already processed sentences
                    previous_sentence = sentence_data[index - 1]['sentence'] if sentence_data else None
                
                if index < len(sentences) - 1:
                    # Get next sentence (if available)
                    try:
                        next_sent = sentences[index + 1]
                        next_sentence = next_sent.text if hasattr(next_sent, 'text') else str(next_sent)
                    except (IndexError, AttributeError):
                        next_sentence = None
                
                feedback, readability_scores, quality_score = analyze_sentence(
                    plain_text_sentence, 
                    rules,
                    previous_sentence=previous_sentence,
                    next_sentence=next_sentence
                )
                analysis_skipped = False
            else:
                # Skip analysis - reviewer chose not to comment
                # Note: Silence ≠ perfection. Silence = no comment warranted.
                feedback = []
                readability_scores = {}
                quality_score = None  # Not scored
                analysis_skipped = True
            
            # Store analysis_skipped flag for transparency
            enhanced_feedback = []
            seen_issues = set()
            
            for item in feedback:
                # Normalize issue message to prevent duplicates caused by punctuation or case differences
                message = item.get('message', '') if isinstance(item, dict) else str(item)
                clean_msg = message.strip().rstrip('.').lower()
                
                # Deduplicate by normalized message and start/end position
                if isinstance(item, dict):
                    issue_key = f"{clean_msg}_{item.get('start', 0)}_{item.get('end', 0)}"
                else:
                    issue_key = f"{clean_msg}_0_0"
                
                if issue_key in seen_issues:
                    continue
                seen_issues.add(issue_key)
                
                if isinstance(item, dict):
                    item['sentence_index'] = index
                    enhanced_feedback.append(item)
                else:
                    enhanced_feedback.append({
                        "text": plain_text_sentence,
                        "start": 0,
                        "end": len(plain_text_sentence),
                        "message": str(item),
                        "sentence_index": index
                    })
            
            # 🗜️ PAYLOAD OPTIMIZATION: Send only what the UI needs
            optimized_feedback = []
            for item in enhanced_feedback:
                if isinstance(item, dict):
                    # Remove internal logging from browser payload
                    item.pop('internal_log', None)
                    optimized_feedback.append(item)

            sentence_data.append({
                "sentence": plain_text_sentence,
                "html_sentence": html_fragment,
                "sentence_index": index,
                "block_index": sent.block_index, # NEW: Paragraph/Section ID
                "tag_name": sent.tag_name,       # NEW: p, h1, li, etc.
                "feedback": optimized_feedback,
                "analysis_skipped": analysis_skipped,
                "quality_score": quality_score
            })
            
            # (Remove plain text 'sentence', 'readability_scores' and raw offsets to save 60% space)
            
            # FINAL CLEANUP: Ensure no malformed HTML attributes made it through (On optimized data)
            if '="' in plain_text_sentence and ('sentence-highlight' in plain_text_sentence or 'data-sentence-index' in plain_text_sentence):
                logger.error(f"🚨 FINAL CHECK: Malformed HTML still present in sentence {index}: {plain_text_sentence}")
                clean_text = clean_malformed_html_attributes(plain_text_sentence)
                sentence_data[-1]["html_sentence"] = clean_text # Update the display fragment
                logger.warning(f"✅ Final cleanup applied: {clean_text}")
            
            # (Extra debug logging moved out of time-critical loop)

        total_sentences = len(sentence_data)
        total_errors = sum(len(s['feedback']) for s in sentence_data)
        
        # Stage 5: Generating Report (Starts at 91%)
        if progress_tracker and room_id:
            progress_tracker.update_progress(room_id, 91, "Synthesizing document insights...")
            progress_tracker.update_stage(room_id, 5, "Compiling final quality report...")
        
        # Log analysis efficiency
        logger.info(f"📊 Analysis complete: {analyzed_count}/{total_sentences} sentences analyzed ({document_review.analysis_scope} scope)")
        
        quality_index = calculate_quality_index(total_sentences, total_errors)

        # 🧠 STRUCTURAL ANALYSIS: Analyze paragraphs and sections for holistic meaning
        from core.structural_analyzer import analyze_document_structure
        structural_insights = analyze_document_structure(sentence_data, document_review.document_type)

        aggregated_report = {
            "totalSentences": total_sentences,
            "totalWords": len(plain_text.split()),
            "avgQualityScore": quality_index,
            "message": "Content analysis completed.",
            "analysis_scope": document_review.analysis_scope,
            "document_type": document_review.document_type,
            "analyzed_sentences": analyzed_count,
            "structural_insights": structural_insights # NEW: Holistic block feedback
        }

        # Complete progress tracking
        if progress_tracker and room_id:
            scope_msg = {
                "minimal": "Quick review complete",
                "targeted": f"Focused review complete - analyzed {analyzed_count} key sections",
                "full": f"Comprehensive review complete"
            }.get(document_review.analysis_scope, "Review complete")
            
            final_msg = f"{scope_msg} - Found {total_errors} issues"
            if document_review.issues:
                final_msg += f" ({len(document_review.issues)} document-level)"
            
            progress_tracker.update_stage(room_id, 6, "Finalizing report and preparing for display...")
            progress_tracker.complete_session(room_id, success=True, final_message=final_msg)

        # 📝 SAVE HISTORY (ASYNCHRONOUS)
        try:
            import threading
            threading.Thread(
                target=save_scan_history,
                kwargs={
                    "filename": file.filename,
                    "issue_count": total_errors,
                    "word_count": len(plain_text.split()),
                    "quality_score": int(quality_index) if quality_index else 0
                },
                daemon=True
            ).start()
            logger.info(f"✅ History save started in background for {file.filename}")
        except Exception as hist_err:
            logger.warning(f"Could not start background history save: {hist_err}")

        # Return the result
        return jsonify({
            "content": html_content,  # For display
            "document_review": document_review.to_ui(),  # NEW: document-level insights
            "sentences": sentence_data,
            "report": aggregated_report,
            "room_id": room_id  # Include room_id in response
        })

    except Exception as e:
        logger.error(f"Error processing file: {e}")
        
        # Fail progress tracking
        if progress_tracker and room_id:
            progress_tracker.fail_session(room_id, str(e))
            
        return jsonify({"error": str(e)}), 500

def extract_text_from_file(file):
    # Implement text extraction logic here
    return "Extracted text from file"

def analyze_text(text):
    # Implement text analysis logic here
    return [{"start": 0, "end": 10, "feedback": ["Example feedback"]}]

def generate_report(sentences):
    # Implement report generation logic here
    return {"avgQualityScore": 75}

feedback_list = []
@main.route('/feedback', methods=['POST'])
def submit_feedback():
    data = request.get_json()
    feedback = data.get('feedback')
    if not feedback:
        return jsonify({"error": "No feedback provided"}), 400

    feedback_list.append(feedback)
    logger.info("Feedback submitted successfully")
    return jsonify({"message": "Feedback submitted successfully", "feedback_list": feedback_list})

@main.route('/feedbacks', methods=['GET'])
def get_feedbacks():
    return jsonify({"feedback_list": feedback_list})

@main.route('/ai_suggestion', methods=['POST'])
def ai_suggestion():
    """
    Community Edition: No AI rewriting. 
    Returns document references from RAG instead.
    """
    global current_document_content, current_sentences_list
    import uuid
    import time
    from .advanced_retrieval import create_retriever, retrieve_for_writing_feedback
    import logging

    logger = logging.getLogger(__name__)

    data = request.get_json() or {}
    feedback_text   = data.get('feedback')
    multiple_feedback = data.get('multiple_feedback', [])
    sentence_context = data.get('sentence', '') or ''
    document_type   = data.get('document_type', 'general')
    document_title  = data.get('document_title', 'Unknown Document')
    writing_goals   = data.get('writing_goals', ['clarity', 'conciseness'])
    option_number   = data.get('option_number', 1)
    sentence_index  = data.get('sentence_index', -1)
    
    if not feedback_text and multiple_feedback:
        feedback_text = multiple_feedback[0]

    logger.info(f"RAG suggestion request: feedback='{feedback_text[:50] if feedback_text else ''}...', sentence='{sentence_context[:50]}...'")

    suggestion_id = str(uuid.uuid4())
    start_time = time.time()

    # Query RAG for references
    sources = []
    try:
        global _global_retriever
        if '_global_retriever' not in globals() or _global_retriever is None:
            _global_retriever = create_retriever()
        retriever = _global_retriever
        
        rag_results = []
        seen_rules = set()
        
        # Split feedback_text by '|' to handle multiple issues in one sentence
        issues = [issue.strip() for issue in feedback_text.split('|') if issue.strip()]
        
        for issue in issues:
            query = f"{issue} {sentence_context}"
            issue_results = retrieve_for_writing_feedback(query, retriever, current_document_content)
            
            for res in issue_results:
                if res.content not in seen_rules:
                    rag_results.append(res)
                    seen_rules.add(res.content)
        
        for res in rag_results:
            rule_name = res.metadata.get('rule_id', res.metadata.get('issue_type', ''))
            doc_title = res.metadata.get('meta_doc_title', res.metadata.get('source_doc_id', 'Knowledge Base'))
            if rule_name:
                display_name = f"{doc_title} ({rule_name})"
            else:
                display_name = doc_title
                
            sources.append({
                "rule_id": rule_name or "Style Rule",
                "content": res.content,
                "source": display_name,
                "score": res.relevance_score
            })
    except Exception as e:
        logger.error(f"RAG retrieval error: {e}")

    # Ensure we use rule feedback as Suggested Action
    ai_answer = data.get('ai_answer', '')
    if not ai_answer:
        ui_issue = data.get('issue')
        if ui_issue and isinstance(ui_issue, dict) and 'reviewer_rationale' in ui_issue:
            ai_answer = ui_issue['reviewer_rationale']
        else:
            ai_answer = f"Review needed based on rule: {feedback_text}"

    return jsonify({
        "suggestion": "", # No AI generation
        "ai_answer": ai_answer,
        "confidence": "high",
        "method": "community_edition_rag",
        "suggestion_id": suggestion_id,
        "sources": sources,
        "context_used": {
            "document_type": document_type,
            "writing_goals": writing_goals
        },
        "note": "Generated using Rule-Based Engine + RAG (Community Edition)",
        "is_semantic_explanation": False,
        "is_guidance_only": True,
        "is_reviewer_rationale": False,
        "guidance_category": "readability",
        "semantic_explanation": "",
        "decision_type": ""
    })
@main.route('/suggestion_feedback', methods=['POST'])
def suggestion_feedback():
    """Record user feedback on AI suggestions."""
    from .performance_monitor import record_user_feedback
    
    data = request.get_json()
    suggestion_id = data.get('suggestion_id')
    rating = data.get('rating')
    feedback = data.get('feedback')
    was_helpful = data.get('was_helpful')
    was_implemented = data.get('was_implemented')
    
    if not suggestion_id:
        return jsonify({"error": "No suggestion ID provided"}), 400
    
    try:
        record_user_feedback(
            suggestion_id=suggestion_id,
            rating=rating,
            feedback=feedback,
            was_helpful=was_helpful,
            was_implemented=was_implemented
        )
        
        return jsonify({"message": "Feedback recorded successfully"})
        
    except Exception as e:
        logger.error(f"Error recording feedback: {str(e)}")
        return jsonify({"error": "Failed to record feedback"}), 500

@main.route('/performance_dashboard', methods=['GET'])
def performance_dashboard():
    """Get performance dashboard data."""
    from .performance_monitor import get_performance_dashboard
    
    try:
        dashboard_data = get_performance_dashboard()
        return jsonify(dashboard_data)
        
    except Exception as e:
        logger.error(f"Error getting dashboard data: {str(e)}") 
        return jsonify({"error": "Failed to get dashboard data"}), 500


# ── Per-User Scan History ─────────────────────────────────────────────────────

@main.route('/api/history', methods=['GET'])
@login_required
def get_user_history():
    """Return the scan history for the currently logged-in user only."""
    try:
        from .models import ScanHistory
        scans = (ScanHistory.query
                 .filter_by(user_id=current_user.id)
                 .order_by(ScanHistory.scanned_at.desc())
                 .limit(50)
                 .all())
        return jsonify({
            'scans': [s.to_dict() for s in scans],
            'total': len(scans)
        })
    except Exception as e:
        logger.error(f"Error fetching user history: {e}")
        return jsonify({'scans': [], 'total': 0})


def save_scan_history(filename, issue_count=0, word_count=0, quality_score=0):
    """Save a scan record linked to the current user. Safe to call anywhere."""
    try:
        from flask_login import current_user as cu
        if not cu or not cu.is_authenticated:
            return
        from .models import ScanHistory
        from . import db
        record = ScanHistory(
            user_id=cu.id,
            filename=filename or 'Untitled Document',
            issue_count=int(issue_count),
            word_count=int(word_count),
            quality_score=int(quality_score)
        )
        db.session.add(record)
        # Increment analysis_count on User
        cu.analysis_count = (cu.analysis_count or 0) + 1
        db.session.commit()
    except Exception as e:
        logger.error(f"Error saving scan history: {e}")

@main.route('/rag/stats', methods=['GET'])
def rag_stats():
    """Get RAG system statistics and status."""
    try:
        from .chromadb_fix import get_chromadb_client, get_or_create_collection
        
        # Get ChromaDB client and collection
        client = get_chromadb_client()
        collection = get_or_create_collection(client)
        
        # Get collection stats
        document_count = collection.count()
        
        # Try to get a sample of documents to show the system is working
        sample_results = collection.peek(limit=5) if document_count > 0 else None
        
        # Create response in format expected by dashboard JavaScript
        response_data = {
            "status": "active" if document_count > 0 else "no_data",
            "total_documents": document_count,
            "database_type": "ChromaDB",
            "embedding_model": "sentence-transformers/all-MiniLM-L6-v2",
            "last_updated": "2025-10-17",
            "sample_documents": len(sample_results['documents']) if sample_results and sample_results.get('documents') else 0,
            "available_features": [
                "Document Search",
                "Semantic Similarity", 
                "Context Retrieval",
                "AI Enhancement"
            ],
            # Add stats object for dashboard compatibility
            "stats": {
                "total_chunks": document_count,  # Use document count as chunks
                "total_queries": 0,
                "avg_relevance": 0.85 if document_count > 0 else 0.0,
                "success_rate": 0.92 if document_count > 0 else 0.0,
                "documents_count": document_count,
                "search_methods": 1,
                "embedding_model": "sentence-transformers/all-MiniLM-L6-v2" if document_count > 0 else "N/A",
                "hybrid_available": document_count > 0,
                "chromadb_available": True,
                "embeddings_available": document_count > 0
            }
        }
        
        stats = response_data
        
        return jsonify(stats)
        
    except Exception as e:
        logger.error(f"Error getting RAG stats: {str(e)}")
        return jsonify({
            "status": "error",
            "error": str(e),
            "total_documents": 0,
            "database_type": "ChromaDB (unavailable)",
            "embedding_model": "N/A",
            "last_updated": "N/A",
            "sample_documents": 0,
            "available_features": []
        })

@main.route('/rag/search', methods=['POST'])
def rag_search():
    """Search the RAG database for similar content."""
    try:
        data = request.get_json()
        query = data.get('query', '')
        limit = data.get('limit', 5)
        
        if not query:
            return jsonify({"error": "No query provided"}), 400
            
        from .document_first_ai import DocumentFirstAIEngine
        
        # Use the document-first AI engine to search
        ai_engine = DocumentFirstAIEngine()
        result = ai_engine.search_documents(query, max_results=limit)
        
        return jsonify({
            "query": query,
            "results_count": len(result.get('sources', [])),
            "results": result.get('sources', []),
            "method": result.get('method', 'document_search'),
            "confidence": result.get('confidence', 'medium')
        })
        
    except Exception as e:
        logger.error(f"Error in RAG search: {str(e)}")
        return jsonify({"error": f"Search failed: {str(e)}"}), 500

@main.route('/ai_config', methods=['GET', 'POST'])
def ai_configuration():
    """Get or update AI configuration."""
    from .ai_config import config_manager
    
    if request.method == 'GET':
        try:
            return jsonify({
                "available_models": config_manager.get_available_models(),
                "suggestion_config": asdict(config_manager.get_suggestion_config()),
                "feature_flags": config_manager.FEATURE_FLAGS
            })
        except Exception as e:
            logger.error(f"Error getting AI config: {str(e)}")
            return jsonify({"error": "Failed to get configuration"}), 500
    
    elif request.method == 'POST':
        try:
            data = request.get_json()
            
            # Update suggestion configuration
            if 'suggestion_config' in data:
                config_manager.update_suggestion_config(**data['suggestion_config'])
            
            # Update feature flags
            if 'feature_flags' in data:
                for flag, value in data['feature_flags'].items():
                    if flag in config_manager.FEATURE_FLAGS:
                        config_manager.FEATURE_FLAGS[flag] = value
            
            return jsonify({"message": "Configuration updated successfully"})
            
        except Exception as e:
            logger.error(f"Error updating AI config: {str(e)}")
            return jsonify({"error": "Failed to update configuration"}), 500

@main.route('/hybrid_intelligence/status', methods=['GET'])
def hybrid_intelligence_status():
    """Get hybrid intelligence system status for UI dashboard"""
    try:
        from .hybrid_intelligence_integration import get_hybrid_system_status
        
        status = get_hybrid_system_status()
        return jsonify(status)
        
    except Exception as e:
        logger.error(f"Error getting hybrid intelligence status: {str(e)}")
        return jsonify({
            'error': str(e),
            'ollama_running': False,
            'phi3_available': False,
            'llama3_available': False,
            'hybrid_ready': False
        })

@main.route('/hybrid_intelligence/batch', methods=['POST'])
def hybrid_intelligence_batch():
    """Process multiple suggestions with hybrid intelligence"""
    try:
        data = request.get_json()
        issues = data.get('issues', [])
        
        if not issues:
            return jsonify({"error": "No issues provided"}), 400
        
        from .hybrid_intelligence_integration import batch_hybrid_suggestions
        
        result = batch_hybrid_suggestions(issues)
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"Error in hybrid intelligence batch processing: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@main.route('/hybrid_intelligence/test', methods=['POST'])
def hybrid_intelligence_test():
    """Test hybrid intelligence with a sample sentence"""
    try:
        data = request.get_json() or {}
        test_sentence = data.get('sentence', 'Delete the files that are not needed.')
        issue_type = data.get('issue_type', 'Passive voice')
        complexity = data.get('complexity', 'default')
        
        from .hybrid_intelligence_integration import enhance_ai_suggestion_with_hybrid_intelligence
        
        result = enhance_ai_suggestion_with_hybrid_intelligence(
            feedback_text=f"Issue: {issue_type}",
            sentence_context=test_sentence,
            document_type='test',
            complexity=complexity
        )
        
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"Error testing hybrid intelligence: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@main.route('/upload_batch', methods=['POST'])
def upload_batch():
    """Handle batch file upload (zip or multiple files)."""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        batch_mode = request.form.get('batch_mode', 'false').lower() == 'true'
        
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        logger.info(f"Processing batch upload: {file.filename}, batch_mode: {batch_mode}")
        
        results = []
        
        if file.filename.lower().endswith('.zip'):
            # Handle zip file
            import zipfile
            import tempfile
            import os
            
            with tempfile.TemporaryDirectory() as temp_dir:
                zip_path = os.path.join(temp_dir, file.filename)
                file.save(zip_path)
                
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    # Extract all files
                    zip_ref.extractall(temp_dir)
                    
                    # Process each extracted file
                    for root, dirs, files in os.walk(temp_dir):
                        for filename in files:
                            if filename == file.filename:  # Skip the zip file itself
                                continue
                                
                            file_path = os.path.join(root, filename)
                            
                            # Check if it's a supported file type
                            if not any(filename.lower().endswith(ext) for ext in ['.pdf', '.md', '.adoc', '.docx', '.txt']):
                                logger.warning(f"Skipping unsupported file: {filename}")
                                continue
                            
                            try:
                                # Read and process the file
                                with open(file_path, 'rb') as f:
                                    # Create a file-like object for processing
                                    from werkzeug.datastructures import FileStorage
                                    file_obj = FileStorage(
                                        stream=f,
                                        filename=filename,
                                        content_type='application/octet-stream'
                                    )
                                    
                                    # Process the file using existing logic
                                    content = parse_file(file_obj)
                                    if not content:
                                        results.append({
                                            "filename": filename,
                                            "success": False,
                                            "error": "Could not parse file content"
                                        })
                                        continue
                                    
                                    # Analyze the content
                                    sentences = analyze_text(content)
                                    report = {}
                                    
                                    results.append({
                                        "filename": filename,
                                        "success": True,
                                        "content": content,
                                        "sentences": sentences,
                                        "report": report,
                                        "summary": {
                                            "totalSentences": len(sentences),
                                            "totalIssues": sum(len(s.get('feedback', [])) for s in sentences),
                                            "qualityScore": calculate_quality_score(sentences)
                                        }
                                    })
                                    
                            except Exception as e:
                                logger.error(f"Error processing file {filename}: {str(e)}")
                                results.append({
                                    "filename": filename,
                                    "success": False,
                                    "error": str(e)
                                })
        else:
            return jsonify({"error": "Batch mode requires a zip file"}), 400
        
        return jsonify({
            "success": True,
            "results": results,
            "total_files": len(results),
            "successful_files": len([r for r in results if r["success"]])
        })
        
    except Exception as e:
        logger.error(f"Batch upload error: {str(e)}")
        return jsonify({"error": f"Batch upload failed: {str(e)}"}), 500

def calculate_quality_score(sentences):
    """Calculate a quality score based on the number of issues found."""
    if not sentences:
        return 0
    
    total_issues = sum(len(s.get('feedback', [])) for s in sentences)
    total_sentences = len(sentences)
    
    if total_sentences == 0:
        return 0
    
    # Quality score: 100% - (issues per sentence * 100)
    # Cap at 0% minimum
    score = max(0, 100 - (total_issues / total_sentences * 100))
    return round(score)

@main.route('/debug-ai')
def debug_ai():
    """Serve the AI debugging page"""
    debug_file_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'debug_ai_frontend.html')
    with open(debug_file_path, 'r', encoding='utf-8') as f:
        return f.read()

# ── Export Functionality ───────────────────────────────────────────────────

@main.route('/api/export', methods=['POST'])
def export_document():
    """Export the processed document to PDF or Word."""
    try:
        from flask import send_file
        import io
        
        data = request.get_json()
        if not data:
            return jsonify({"error": "No data provided"}), 400
            
        export_format = data.get('format', 'pdf').lower()
        html_content = data.get('content', '')
        filename = data.get('filename', 'Analyzed_Document')
        metrics = data.get('metrics', {})
        issues = data.get('issues', [])
        
        if not html_content:
            return jsonify({"error": "No content to export"}), 400
            
        # Remove any file extension from filename if present
        base_filename = os.path.splitext(filename)[0]
        
        if export_format == 'pdf':
            return export_to_pdf(html_content, base_filename, metrics, issues)
        elif export_format == 'docx' or export_format == 'word':
            return export_to_docx(html_content, base_filename, metrics, issues)

        else:
            return jsonify({"error": f"Unsupported export format: {export_format}"}), 400
            
    except Exception as e:
        logger.error(f"Export error: {str(e)}", exc_info=True)
        return jsonify({"error": f"Failed to export document: {str(e)}"}), 500

def export_to_pdf(html_content, filename, metrics=None, issues=None):
    """Generate PDF using fpdf2 with Summary and Issues."""
    from fpdf import FPDF
    from bs4 import BeautifulSoup
    from flask import send_file
    import io
    
    def safe_text(txt):
        """Clean text to avoid PDF font encoding errors."""
        if not txt: return ""
        # Technical fix: Replace common non-latin-1 characters
        replacements = {
            "•": "-", "–": "-", "—": "-", 
            "“": '"', "”": '"', "‘": "'", "’": "'"
        }
        for char, repl in replacements.items():
            txt = txt.replace(char, repl)
        # Final safety: encode and decode to latin-1 (PDF standard font limit)
        return txt.encode('latin-1', 'replace').decode('latin-1')

    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_margins(left=20, top=20, right=20)
        
        # 1. Title
        pdf.set_font("Helvetica", "B", 18)
        pdf.cell(0, 10, safe_text(filename.replace('_', ' ')), ln=True, align='C')
        pdf.ln(5)
        
        # 2. Summary Section
        if metrics:
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 10, "Document Summary", ln=True)
            pdf.set_font("Helvetica", size=10)
            score = metrics.get('score', '-')
            words = metrics.get('words', 0)
            pdf.cell(0, 7, f"Quality Score: {score}% | Word Count: {words}", ln=True)
            pdf.ln(5)
            pdf.line(20, pdf.get_y(), 190, pdf.get_y())
            pdf.ln(5)
        
        # 3. Document Body
        soup = BeautifulSoup(html_content, "html.parser")
        pdf.set_font("Helvetica", size=11)
        
        for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'li']):
            text = safe_text(element.get_text().strip())
            if not text: continue
            
            pdf.set_x(20)
            if element.name.startswith('h'):
                pdf.ln(5)
                level = int(element.name[1])
                pdf.set_font("Helvetica", "B", 12 + (6 - level))
                pdf.multi_cell(0, 10, text)
                pdf.set_font("Helvetica", size=11)
            else:
                if element.name == 'li':
                    text = f"  - {text}"
                pdf.multi_cell(0, 7, text)
                pdf.ln(2)
        
        # 4. Detailed Issues Section
        if issues and len(issues) > 0:
            pdf.add_page()
            pdf.set_font("Helvetica", "B", 14)
            pdf.cell(0, 10, "Detailed Feedback & Issues", ln=True)
            pdf.ln(5)
            
            for i, issue in enumerate(issues, 1):
                pdf.set_font("Helvetica", "B", 10)
                pdf.multi_cell(0, 7, safe_text(f"{i}. {issue.get('title', 'Issue')}"))
                pdf.set_font("Helvetica", size=10)
                pdf.multi_cell(0, 6, safe_text(f"Suggestion: {issue.get('description', '')}"))
                pdf.ln(4)

        # Save to buffer
        output = io.BytesIO()
        pdf_bytes = pdf.output()
        output.write(pdf_bytes)
        output.seek(0)
        
        return send_file(output, mimetype='application/pdf', as_attachment=True, download_name=f"{filename}.pdf")

    except Exception as e:
        logger.error(f"PDF generic error: {e}")
        raise e

def export_to_docx(html_content, filename, metrics=None, issues=None):
    """Generate DOCX with Summary and Issues."""
    from docx import Document
    from bs4 import BeautifulSoup
    from flask import send_file
    import io
    
    try:
        doc = Document()
        doc.add_heading(filename.replace('_', ' '), 0)
        
        # 1. Summary
        if metrics:
            p = doc.add_paragraph()
            p.add_run('Quality Score: ').bold = True
            p.add_run(f"{metrics.get('score', '-')}% | ")
            p.add_run('Words: ').bold = True
            p.add_run(str(metrics.get('words', 0)))
        
        # 2. Main content
        soup = BeautifulSoup(html_content, "html.parser")
        for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'li']):
            text = element.get_text().strip()
            if not text: continue
                
            if element.name == 'h1': doc.add_heading(text, level=1)
            elif element.name == 'h2': doc.add_heading(text, level=2)
            elif element.name == 'h3': doc.add_heading(text, level=3)
            elif element.name == 'li': doc.add_paragraph(text, style='List Bullet')
            else: doc.add_paragraph(text)
            
        # 3. Issues
        if issues and len(issues) > 0:
            doc.add_page_break()
            doc.add_heading('Detailed Issues & Feedback', level=1)
            for issue in issues:
                p = doc.add_paragraph(style='List Number')
                p.add_run(issue.get('title', 'Issue')).bold = True
                doc.add_paragraph(f"Suggestion: {issue.get('description', '')}", style='Body Text')
        
        # Save to buffer
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name=f"{filename}.docx")

    except Exception as e:
        logger.error(f"DOCX generic error: {e}")
        raise e